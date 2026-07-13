from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FINAL_DELIVERY.md"
OUTPUT = ROOT / "docs" / "智能相册分析系统-完整交付文档.docx"
TEAL = "0F766E"
LIGHT_TEAL = "F0FDFA"
AMBER = "D97706"
INK = "17201F"
MUTED = "64706E"
LINE = "D9E5E2"
CJK_FONT = "STHeiti"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)


def set_run_font(run, name=CJK_FONT, size=None, color=None, bold=None):
    run.font.name = name
    properties = run._element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{slot}"), name)
    fonts.set(qn("w:hint"), "eastAsia")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=8.5, color=MUTED)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = CJK_FONT
    normal_fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal_fonts.set(qn(f"w:{slot}"), CJK_FONT)
    normal_fonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    style_specs = {
        "Title": (27, TEAL, True, 0, 16),
        "Heading 1": (16, TEAL, True, 16, 7),
        "Heading 2": (12.5, INK, True, 12, 5),
        "Heading 3": (10.5, TEAL, True, 9, 4),
    }
    for style_name, (size, color, bold, before, after) in style_specs.items():
        style = document.styles[style_name]
        style.font.name = CJK_FONT
        style_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            style_fonts.set(qn(f"w:{slot}"), CJK_FONT)
        style_fonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "智能相册分析系统 | 完整交付文档"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def strip_inline_markdown(text):
    text = re.sub(r"\[([^]]+)]\(([^)]+)\)", r"\1（\2）", text)
    text = re.sub(r"<((?:https?|file)://[^>]+)>", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return text


def add_rich_text(paragraph, text):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|<https?://[^>]+>|\[[^]]+\]\([^)]+\))")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Menlo", size=8.5, color=TEAL)
        elif token.startswith("<"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, color=TEAL)
            run.underline = True
        else:
            label, url = re.match(r"\[([^]]+)]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label}（{url}）")
            set_run_font(run, color=TEAL)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run)


def add_code_block(document, lines):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F7F6")
    properties.append(shading)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, name="Menlo", size=7.8, color=INK)


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows


def add_table(document, rows):
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            value = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_rich_text(paragraph, value)
            for run in paragraph.runs:
                set_run_font(
                    run,
                    size=8.2,
                    color="FFFFFF" if row_index == 0 else INK,
                    bold=row_index == 0,
                )
            if row_index == 0:
                set_cell_shading(cell, TEAL)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F7FAF9")
    set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image(document, image_path, alt_text):
    if not image_path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.45))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    caption_run = caption.add_run(f"图 1  {alt_text}")
    set_run_font(caption_run, size=8.5, color=MUTED)


def add_cover(document, title, metadata):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(46)
    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("真实照片 · 可解释分析 · 在线与离线双交付")
    set_run_font(run, size=12, color=AMBER, bold=True)

    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line.add_run("━" * 28)
    set_run_font(line_run, size=8, color=TEAL)

    summary = document.add_table(rows=1, cols=1)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    summary.columns[0].width = Inches(5.8)
    cell = summary.cell(0, 0)
    set_cell_shading(cell, LIGHT_TEAL)
    set_cell_margins(cell, top=220, start=260, bottom=220, end=260)
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_paragraph.add_run("76 张脱敏真实照片 · 11 项自动化测试 · 自包含离线演示")
    set_run_font(run, size=11, color=TEAL, bold=True)

    document.add_paragraph().paragraph_format.space_after = Pt(18)
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, size=9.5, color=MUTED, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=9.5, color=INK)

    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.space_before = Pt(22)
    status_run = status.add_run("当前状态：代码、数据、演示、测试、文档与答辩 PPT 已齐备")
    set_run_font(status_run, size=9.5, color=TEAL, bold=True)
    status.add_run().add_break(WD_BREAK.PAGE)


def build_document():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    metadata = []
    cursor = 1
    while cursor < len(lines):
        line = lines[cursor].strip()
        match = re.match(r"\*\*(.+?)：\*\*\s*(.+?)(?:\s{2})?$", line)
        if match:
            metadata.append((match.group(1), strip_inline_markdown(match.group(2))))
        elif line.startswith("## "):
            break
        cursor += 1

    document = Document()
    configure_document(document)
    add_cover(document, title, metadata)

    in_code = False
    code_lines = []
    index = cursor
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("| "):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, parse_table(table_lines))
            continue
        image_match = re.fullmatch(r"!\[([^]]*)]\(([^)]+)\)", stripped)
        if image_match:
            add_image(document, (SOURCE.parent / image_match.group(2)).resolve(), image_match.group(1))
            index += 1
            continue
        if stripped.startswith("## "):
            document.add_paragraph(stripped[3:], style="Heading 1")
        elif stripped.startswith("### "):
            document.add_paragraph(stripped[4:], style="Heading 2")
        elif stripped.startswith("#### "):
            document.add_paragraph(stripped[5:], style="Heading 3")
        elif re.match(r"^- \[[ xX]] ", stripped):
            checked = stripped[3].lower() == "x"
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, f"{'已完成' if checked else '待完成'}：{stripped[6:]}")
            paragraph.paragraph_format.left_indent = Inches(0.18)
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, stripped[2:])
            paragraph.paragraph_format.left_indent = Inches(0.18)
        elif re.match(r"^\d+\. ", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_rich_text(paragraph, re.sub(r"^\d+\. ", "", stripped))
            paragraph.paragraph_format.left_indent = Inches(0.18)
        elif stripped.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            properties = paragraph._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), LIGHT_TEAL)
            properties.append(shading)
            add_rich_text(paragraph, stripped[2:])
        else:
            paragraph = document.add_paragraph()
            add_rich_text(paragraph, stripped)
        index += 1

    core = document.core_properties
    core.title = title
    core.subject = "课程项目完整交付说明"
    core.author = "智能相册分析系统项目组"
    core.keywords = "智能相册, 数据分析, React, Flask, ECharts"
    core.comments = "由项目源码中的可复现脚本生成"
    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    build_document()
